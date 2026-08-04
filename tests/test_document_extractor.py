# Synthelion — Python port of Caveman (https://github.com/francescopaolopassaro/caveman)
# © 2026 Passaro Francesco Paolo — Digitalsolutions.it
"""Per-format document extraction/masking tests for `synthelion/document_extractor.py`.

DOCX/XLSX in-place masking tests are skipped when python-docx/openpyxl aren't
installed (optional `synthelion[documents]` extra) — same convention as the
existing HAS_CREWAI/HAS_LANGCHAIN skip pattern elsewhere in this suite. PDF
extraction is tested by faking `pypdf` in `sys.modules` so the chunking logic
is covered without requiring the real dependency.
"""
from __future__ import annotations

import csv
import sys
import types
from unittest.mock import patch

import pytest

from synthelion import document_extractor as de
from synthelion.privacy_analyzer import PrivacyAnalyzer
from synthelion.privacy_session import PrivacySession

try:
    import docx  # noqa: F401
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl  # noqa: F401
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False


class TestDetectFormat:
    def test_known_extensions(self):
        assert de.detect_format("a.pdf") == "pdf"
        assert de.detect_format("a.docx") == "docx"
        assert de.detect_format("a.xlsx") == "xlsx"
        assert de.detect_format("a.csv") == "csv"
        assert de.detect_format("a.md") == "md"
        assert de.detect_format("a.markdown") == "md"

    def test_unknown_extension_falls_back_to_txt(self):
        assert de.detect_format("a.unknown") == "txt"
        assert de.detect_format("a.txt") == "txt"


class TestPdfExtraction:
    def test_extract_pdf_text_chunks(self):
        fake_page1 = types.SimpleNamespace(extract_text=lambda: "Hello world. " * 5)
        fake_page2 = types.SimpleNamespace(extract_text=lambda: "Second page content.")
        fake_reader = types.SimpleNamespace(pages=[fake_page1, fake_page2])
        fake_pypdf = types.SimpleNamespace(PdfReader=lambda path: fake_reader)

        with patch.dict(sys.modules, {"pypdf": fake_pypdf}):
            chunks = list(de.extract_pdf_text("fake.pdf", chunk_chars=20))

        joined = "".join(chunks)
        assert "Hello world." in joined
        assert "Second page content." in joined
        for chunk in chunks[:-1]:
            assert len(chunk) == 20

    def test_missing_dependency_raises_clear_error(self):
        with patch.dict(sys.modules, {"pypdf": None}):
            with pytest.raises(ImportError, match=r"synthelion\[documents\]"):
                list(de.extract_pdf_text("fake.pdf"))


class TestCsvExtraction:
    def test_extract_csv_rows_streaming(self, tmp_path):
        p = tmp_path / "data.csv"
        with open(p, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(["name", "email"])
            for i in range(5):
                w.writerow([f"user{i}", f"user{i}@example.com"])

        chunks = list(de.extract_csv_rows(str(p), chunk_rows=2))
        joined = "".join(chunks)
        assert "user0@example.com" in joined
        assert "user4@example.com" in joined
        assert len(chunks) == 3  # ceil(6 rows / 2)


class TestMarkdownExtraction:
    def test_extract_text_chunks_markdown(self, tmp_path):
        p = tmp_path / "doc.md"
        content = "# Title\n\n" + ("word " * 5000)
        p.write_text(content, encoding="utf-8")
        chunks = list(de.extract_text_chunks(str(p), chunk_chars=1024))
        assert "".join(chunks) == content
        assert len(chunks) > 1


class TestPlainTextExtraction:
    def test_extract_text_chunks_plain(self, tmp_path):
        p = tmp_path / "doc.txt"
        p.write_text("Contact us at test@example.com for help.", encoding="utf-8")
        chunks = list(de.extract_text_chunks(str(p)))
        assert "".join(chunks) == "Contact us at test@example.com for help."


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed (pip install synthelion[documents])")
class TestDocxInPlaceMasking:
    def test_mask_docx_in_place(self, tmp_path):
        src = tmp_path / "sample.docx"
        d = docx.Document()
        d.add_paragraph("Please contact John at john.doe@example.com for details.")
        d.add_paragraph("This paragraph has no sensitive data at all.")
        table = d.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "IBAN: IT60X0542811101000000123456"
        d.save(str(src))

        analyzer = PrivacyAnalyzer()
        session = PrivacySession()
        doc = de.load_docx(str(src))
        count = de.mask_docx_in_place(doc, analyzer, session, "en")
        out = tmp_path / "sample.masked.docx"
        doc.save(str(out))

        assert count > 0
        reopened = docx.Document(str(out))
        full_text = "\n".join(p.text for p in reopened.paragraphs)
        full_text += "\n" + reopened.tables[0].rows[0].cells[0].text
        assert "john.doe@example.com" not in full_text
        assert "This paragraph has no sensitive data at all." in full_text


@pytest.mark.skipif(not HAS_OPENPYXL, reason="openpyxl not installed (pip install synthelion[documents])")
class TestXlsxInPlaceMasking:
    def test_mask_xlsx_in_place(self, tmp_path):
        src = tmp_path / "sample.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"] = "Name"
        ws["B1"] = "Email"
        ws["A2"] = "Jane"
        ws["B2"] = "jane.smith@example.com"
        wb.save(str(src))

        analyzer = PrivacyAnalyzer()
        session = PrivacySession()
        wb2 = de.load_xlsx(str(src))
        count = de.mask_xlsx_in_place(wb2, analyzer, session, "en")
        out = tmp_path / "sample.masked.xlsx"
        wb2.save(str(out))

        assert count > 0
        reopened = openpyxl.load_workbook(str(out))
        ws2 = reopened.active
        assert ws2["B2"].value != "jane.smith@example.com"
        assert ws2["A2"].value == "Jane"
