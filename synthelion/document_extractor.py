# Synthelion — Python port of Caveman (https://github.com/francescopaolopassaro/caveman)
# © 2026 Passaro Francesco Paolo — Digitalsolutions.it
"""Per-format document extraction and in-place PII masking for PrivacyGuard.

PDF/CSV/Markdown/plain-text are read in chunks (never the whole file in RAM).
DOCX/XLSX in-place masking necessarily loads the full python-docx/openpyxl
object model — those libraries expose no incremental write API for in-place
edits, so "streaming" for those two formats means the read side of things,
not the write side.

All PDF/DOCX/XLSX support requires the optional `synthelion[documents]` extra
(`pypdf`, `python-docx`, `openpyxl`). CSV/Markdown/plain-text need nothing
beyond the stdlib.
"""
from __future__ import annotations

import csv
import pathlib
from typing import TYPE_CHECKING, Iterator

if TYPE_CHECKING:
    from synthelion.privacy_analyzer import PrivacyAnalyzer
    from synthelion.privacy_session import PrivacySession

_FORMAT_BY_EXT = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".xlsx": "xlsx",
    ".csv": "csv",
    ".md": "md",
    ".markdown": "md",
}


def detect_format(path: str) -> str:
    """Extension-based format dispatch. Unrecognized extensions fall back to
    plain text (".txt" behavior) rather than raising."""
    ext = pathlib.Path(path).suffix.lower()
    return _FORMAT_BY_EXT.get(ext, "txt")


def _missing_dependency(package: str, pip_extra: str = "documents") -> ImportError:
    return ImportError(
        f"'{package}' is required for this document format but is not installed. "
        f"Install it with: pip install \"synthelion[{pip_extra}]\""
    )


# ── PDF ──────────────────────────────────────────────────────────────────────

def extract_pdf_text(path: str, *, chunk_chars: int = 8192) -> Iterator[str]:
    """Extract text from a PDF page-by-page, yielded in ``chunk_chars``-sized
    pieces so a large PDF's full text is never held in memory at once."""
    try:
        import pypdf
    except ImportError:
        raise _missing_dependency("pypdf") from None

    reader = pypdf.PdfReader(path)
    buffer = ""
    for page in reader.pages:
        buffer += (page.extract_text() or "") + "\n"
        while len(buffer) >= chunk_chars:
            yield buffer[:chunk_chars]
            buffer = buffer[chunk_chars:]
    if buffer:
        yield buffer


# ── DOCX ─────────────────────────────────────────────────────────────────────

def load_docx(path: str):
    try:
        import docx
    except ImportError:
        raise _missing_dependency("python-docx") from None
    return docx.Document(path)


def _mask_paragraphs(paragraphs, analyzer: "PrivacyAnalyzer", session: "PrivacySession", language: str) -> int:
    masked_count = 0
    for para in paragraphs:
        original = para.text
        if not original or not original.strip():
            continue
        result = analyzer.analyze(original, language=language, session=session, auto_masking=True)
        if result.masked_text and result.masked_text != original:
            masked_count += result.match_count
            # Masking collapses a paragraph's runs into the first run's formatting
            # (mixed-formatting runs within a single masked paragraph are not
            # individually preserved) — paragraphs with no PII are left untouched.
            if para.runs:
                para.runs[0].text = result.masked_text
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.add_run(result.masked_text)
    return masked_count


def mask_docx_in_place(doc, analyzer: "PrivacyAnalyzer", session: "PrivacySession", language: str = "en") -> int:
    """Masks PII in every paragraph, table cell, and header/footer of *doc*
    (a `docx.Document`), replacing text directly. Returns total match count."""
    masked = _mask_paragraphs(doc.paragraphs, analyzer, session, language)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                masked += _mask_paragraphs(cell.paragraphs, analyzer, session, language)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            masked += _mask_paragraphs(hf.paragraphs, analyzer, session, language)
    return masked


# ── XLSX ─────────────────────────────────────────────────────────────────────

def load_xlsx(path: str):
    try:
        import openpyxl
    except ImportError:
        raise _missing_dependency("openpyxl") from None
    return openpyxl.load_workbook(path)


def mask_xlsx_in_place(wb, analyzer: "PrivacyAnalyzer", session: "PrivacySession", language: str = "en") -> int:
    """Masks PII in every string cell of every worksheet of *wb* (an openpyxl
    `Workbook`), writing back `cell.value`. Returns total match count."""
    masked_count = 0
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                value = cell.value
                if not isinstance(value, str) or not value.strip():
                    continue
                result = analyzer.analyze(value, language=language, session=session, auto_masking=True)
                if result.masked_text and result.masked_text != value:
                    masked_count += result.match_count
                    cell.value = result.masked_text
    return masked_count


# ── CSV / Markdown / plain text ─────────────────────────────────────────────

def extract_csv_rows(path: str, *, chunk_rows: int = 500, encoding: str = "utf-8") -> Iterator[str]:
    """stdlib `csv` streaming reader, re-joined as comma-separated text per
    chunk of ``chunk_rows`` rows."""
    with open(path, "r", encoding=encoding, newline="") as f:
        reader = csv.reader(f)
        buffer: list[str] = []
        for row in reader:
            buffer.append(",".join(row))
            if len(buffer) >= chunk_rows:
                yield "\n".join(buffer) + "\n"
                buffer = []
        if buffer:
            yield "\n".join(buffer) + "\n"


def extract_text_chunks(path: str, *, chunk_chars: int = 8192, encoding: str = "utf-8") -> Iterator[str]:
    """Plain chunked read for Markdown/plain-text/any unrecognized format —
    a large file is never loaded whole."""
    with open(path, "r", encoding=encoding, errors="replace") as f:
        while True:
            chunk = f.read(chunk_chars)
            if not chunk:
                break
            yield chunk
